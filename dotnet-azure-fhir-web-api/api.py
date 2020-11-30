from flask import Flask, jsonify, render_template, request, redirect, url_for, make_response, send_file
from flask_restful import Api, Resource
from flask_table import Table, Col
from fhir_parser.fhir import FHIR
import pdfkit
from io import StringIO
from docx import Document



app = Flask(__name__)
api = Api(app)

def _init_(self):
    super(Data, self)._init_()

@app.route("/createRecord/<id>/", methods=['POST'])
def patient_record(id):
    fhir_parser = FHIR(endpoint="http://localhost:5000/api/", verify_ssl=False)
    patient = fhir_parser.get_patient(id)
    name = patient.name.full_name
    gender = patient.gender
    telecom = patient.telecoms
    language = str(patient.communications.languages)
    address = str(patient.addresses)
    birthdate = str(patient.birth_date)
    marital = patient.marital_status
    return render_template("patient.html", id = id, name = name, gender = gender, telecom = telecom, birthdate = birthdate, marital = marital, language = language)

@app.route("/patientPDF/<id>/", methods=['POST'])
def patient_pdf(id):
    template = patient_record(id)
    pdf = pdfkit.from_string(template, False)
    response = make_response(pdf)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'Attachment; filename=patientRecord.pdf'
    return response

@app.route("/patientDOCX/<id>/", methods=['POST'])
def patient_docx(id):
    doc = Document()

    fhir_parser = FHIR(endpoint="http://localhost:5000/api/", verify_ssl=False)
    patient = fhir_parser.get_patient(id)

    doc.add_heading("Patient record")
    doc.add_paragraph("") 
    par = doc.add_paragraph("Patient ID       : ")
    par.add_run(id)
    par = doc.add_paragraph("Full name       : ")
    par.add_run(patient.name.full_name)
    par = doc.add_paragraph("Gender           : ")
    par.add_run(patient.gender)
    par = doc.add_paragraph("Language       : ")
    par.add_run(str(patient.communications.languages))
    par = doc.add_paragraph("Birthdate        : ")
    par.add_run(str(patient.birth_date))
    par = doc.add_paragraph("Marital status : ")
    par.add_run(str(patient.marital_status))

    doc.save("patientRecord.docx")
    return render_template("filecreated.html")


@app.route("/home")   
def make_list():
    patientID_list = []
    patientFullName_list = []
    patientdate_list = []

    fhir_parser = FHIR(endpoint="http://localhost:5000/api/", verify_ssl=False)

    for patient in fhir_parser.get_all_patients():
        id = patient.uuid
        name = patient.name.full_name
        date = str(patient.birth_date)

        patientID_list.append(id)
        patientFullName_list.append(name)
        patientdate_list.append(date)
    
    return render_template("home.html", table_info = zip(patientID_list, patientFullName_list,patientdate_list))


#api.add_resource(Data, "/p/eunice/", endpoint="letter")
if __name__ == "__main__":
    app.run(debug=True, port=5002)



#References
#https://python-docx.readthedocs.io/en/latest/
#https://www.youtube.com/watch?v=C8jxInLM9nM
#https://www.w3schools.com/html/

